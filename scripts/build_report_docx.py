#!/usr/bin/env python3
"""Generate the MetricLens AI development report in USENIX .docx format.

Bilingual: English is primary (development_report_usenix.docx, English figures);
Korean is development_report_usenix_kr.docx (Korean _kr figures, Malgun Gothic
East-Asian font). Embeds the architecture/runtime/approach diagrams and the five
evaluation charts.

Usage:
    python scripts/evaluate_model_paper.py        # figures first (EN + _kr)
    python scripts/build_architecture_diagram.py
    python scripts/build_runtime_diagram.py
    python scripts/build_approach_diagram.py
    python scripts/build_report_docx.py [FRONTEND_URL] [BACKEND_URL]
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "usenix2022.docx"
DIAG = ROOT / "docs" / "diagrams"
EVAL = ROOT / "docs" / "evaluation"

FE = sys.argv[1] if len(sys.argv) > 1 else "(deployed Cloud Run URL)"
BE = sys.argv[2] if len(sys.argv) > 2 else "(deployed Cloud Run URL)"


def fig(name, lang):
    """Figure path: English primary (base), Korean (_kr)."""
    suf = "_kr" if lang == "kr" else ""
    base = DIAG if name in ("architecture", "runtime", "approach", "method_math") else EVAL
    return base / f"{name}{suf}.png"


# Section content: each item is (kind, en, kr). kind in H1/H2/B/IMG/URL.
# IMG payload is (figure_name, en_caption, kr_caption, width_inches).
def sections(lang):
    return [
        ("TITLE",
         "MetricLens AI: Lightweight Time-Series Forecasting and Integer-"
         "Programming Resizing for On-Premise Server Optimization",
         "MetricLens AI: 온프레미스 서버 최적화를 위한 경량 시계열 예측과 정수계획 기반 리사이징"),
        ("AUTH", "Lee Yongwon (lead), Seo Hannyeong, Song Simwoo",
         "이용원 (팀장), 서한녕, 송심우"),
        ("AFFIL", "Team Googling — googai_congress", "Team 구글링 — googai_congress"),

        ("H1", "Abstract", "초록 (Abstract)"),
        ("B",
         "Over-provisioning dominates on-premise infrastructure, wasting OPEX and "
         "energy. MetricLens AI is a self-contained web platform that ingests "
         "multi-dimensional server metrics (CPU, memory, network), forecasts load "
         "with a GPU-free, CPU-only time-series model, and derives SLO-aware minimal "
         "allocations via integer programming. The forecaster combines seasonal-trend "
         "decomposition with an AR(1) residual correction and sizes its prediction "
         "interval from out-of-sample backtest error. The resizing engine solves an "
         "exact integer program over a bounded space. The system is deployed GCP-"
         "natively on Cloud Run via Cloud Build, ingests real Compute Engine instance "
         "metrics from Cloud Monitoring, and resizes real VMs within a cost budget. On "
         "representative data grounded in public datacentre traces (Azure, Alibaba), the "
         "model beats the seasonal-naive baseline (RMSE) on all six workload archetypes "
         "(both baselines on 5/6; Diebold-Mariano p<0.05 on 4/6), with 95% prediction-"
         "interval coverage of 0.93-0.98.",
         "온프레미스 인프라는 과프로비저닝이 만연해 운영비와 전력을 낭비한다. MetricLens AI는 "
         "다차원 서버 메트릭(CPU·메모리·네트워크)을 적재하고, GPU 없이 CPU만으로 구동되는 경량 "
         "시계열 모델로 부하를 예측하며, 정수계획법으로 SLO를 보장하는 최소 자원을 산출하는 자립형 "
         "웹 플랫폼이다. 예측기는 계절-추세 분해에 AR(1) 잔차 보정을 결합하고, 예측 구간을 표본 외 "
         "백테스트 오차로 산정한다. 리사이징 엔진은 유계 공간에서 정확한 정수계획을 푼다. 시스템은 "
         "Cloud Build로 Cloud Run에 GCP 네이티브 배포되며, Cloud Monitoring으로 실제 Compute "
         "Engine 인스턴스 메트릭을 적재하고 예산 한도 내에서 실제 VM을 리사이즈한다. 공개 데이터센터 "
         "트레이스(Azure·Alibaba)에 근거한 대표 데이터에서, 모델은 6개 워크로드 아키타입 모두에서 "
         "seasonal-naive 기준선을 RMSE로 능가하고(5/6에서 두 기준선 모두, Diebold-Mariano p<0.05 "
         "4/6), 95% 예측구간 커버리지는 0.93–0.98로 잘 보정된다."),

        ("H1", "1. Introduction", "1. 서론"),
        ("B",
         "Conservative capacity planning leaves servers chronically idle (the Azure "
         "trace shows 60% of VMs average below 20% CPU). Public-cloud optimizers are "
         "unavailable in air-gapped or regulated on-premise environments. MetricLens AI "
         "targets this gap: an external-dependency-free, GPU-free system that turns "
         "telemetry into quantitative right-sizing guidance. Contributions: (i) a "
         "standard-library seasonal-trend + AR forecaster; (ii) an exact integer-"
         "programming resizing engine with a percentile-peak safety statistic; (iii) "
         "Cloud Monitoring-based real-instance monitoring and budget-guarded real "
         "resize; (iv) representative data grounded in public traces with a paper-level "
         "evaluation.",
         "보수적 용량 산정은 서버를 만성적으로 유휴 상태로 둔다(Azure 트레이스: VM의 60%가 평균 CPU "
         "20% 미만). 퍼블릭 클라우드 최적화 도구는 망분리·규제 온프레미스에서 쓸 수 없다. MetricLens "
         "AI는 이 공백을 겨냥한 외부 의존성 없는 GPU-프리 시스템으로 텔레메트리를 정량적 라이트사이징 "
         "지침으로 전환한다. 기여: (i) 표준 라이브러리 계절-추세+AR 예측기; (ii) 백분위 피크 안전 "
         "통계를 갖춘 정확한 정수계획 리사이징 엔진; (iii) Cloud Monitoring 기반 실제 인스턴스 "
         "모니터링과 예산 가드 실제 리사이즈; (iv) 공개 트레이스 근거 대표 데이터와 논문 수준 평가."),

        ("H1", "2. System Architecture", "2. 시스템 아키텍처"),
        ("B",
         "The architecture is organised as three logical layers, with the runtime "
         "platform and CI/CD modelled as a separate cross-cutting tier rather than a "
         "fourth layer. The Presentation layer is a React 19 SPA served by an "
         "unprivileged nginx container; the Business layer is a single FastAPI app "
         "internally split into Controller, Service, and Repository, with the "
         "forecasting/optimization core isolated as pure, dependency-free code; and the "
         "Data layer is PostgreSQL (Cloud SQL) in production or embedded SQLite for the "
         "demo. Dependencies flow strictly downward, and the pure core is unit-tested "
         "without a database.",
         "아키텍처는 3개 논리 계층으로 구성되며, 런타임 플랫폼과 CI/CD는 4번째 계층이 아니라 별도의 "
         "교차 관심사(cross-cutting) 계층으로 모델링한다. 표현(Presentation) 계층은 비특권 nginx가 "
         "서빙하는 React 19 SPA, 업무(Business) 계층은 Controller·Service·Repository로 분리된 단일 "
         "FastAPI이며 예측·최적화 Core는 순수 코드로 격리된다. 데이터(Data) 계층은 프로덕션 "
         "PostgreSQL(Cloud SQL), 데모 내장 SQLite이다. 의존성은 아래로만 흐르고 순수 Core는 DB 없이 "
         "단위 테스트된다."),
        ("IMG", ("architecture",
                 "Figure 1. Three logical layers; runtime platform & CI/CD as a separate "
                 "cross-cutting tier (official OSS brand logos).",
                 "그림 1. 3개 논리 계층, 런타임 플랫폼·CI/CD는 별도 교차 계층 (공식 OSS 브랜드 로고).", 6.4)),
        ("H2", "2.1. How it runs (operational architecture)", "2.1. 구동 방식 (어떤 아키텍처로 운영되는가)"),
        ("B",
         "MetricLens runs GCP-natively and serverless. Front and back ends are built into "
         "container images, pushed to Artifact Registry, and deployed to Cloud Run "
         "(scaling to zero when idle). A Cloud Build pipeline runs lint -> test -> build "
         "-> deploy -> health check. Using its runtime service account, the backend reads "
         "labelled real-instance CPU (agentless) and memory (Ops Agent) from Cloud "
         "Monitoring, and changes real machine types via the Compute Engine API "
         "(stop -> setMachineType -> start). Secrets come from Secret Manager, never "
         "hard-coded.",
         "MetricLens는 GCP 네이티브 서버리스로 구동된다. 프론트·백엔드는 컨테이너 이미지로 빌드되어 "
         "Artifact Registry에 푸시되고 Cloud Run으로 배포된다(유휴 시 0으로 스케일). Cloud Build "
         "파이프라인이 린트→테스트→빌드→배포→헬스 체크를 수행한다. 백엔드는 런타임 서비스 계정으로 "
         "라벨된 실제 인스턴스의 CPU(에이전트 없이)와 메모리(Ops Agent)를 Cloud Monitoring에서 읽고, "
         "Compute Engine API로 실제 머신 타입을 변경(stop→setMachineType→start)한다. 비밀값은 "
         "Secret Manager에서 주입되며 하드코딩되지 않는다."),
        ("IMG", ("runtime",
                 "Figure 2. Runtime & CI/CD operation: Cloud Build -> Cloud Run, "
                 "Cloud Monitoring ingest + Compute Engine resize.",
                 "그림 2. 런타임·CI/CD 구동 방식: Cloud Build → Cloud Run, "
                 "Cloud Monitoring 실측 적재·Compute Engine 실제 리사이즈.", 6.4)),

        ("H1", "3. Lightweight Forecasting Engine", "3. 경량 예측 엔진"),
        ("IMG", ("approach",
                 "Figure 3. The MetricLens method: telemetry -> decomposition+AR forecast "
                 "(95% interval) -> p95 peak -> integer program -> machine type -> resize.",
                 "그림 3. MetricLens 접근법(방법) 개요: 텔레메트리 → 분해+AR 예측(95% 구간) → p95 피크 "
                 "→ SLO 제약 정수계획 → GCP 머신 타입 → 리사이즈.", 6.6)),
        ("H2", "3.1. End-to-end computation (the equations the code evaluates)",
         "3.1. 전 과정 연산 (코드가 계산하는 수식)"),
        ("B",
         "The full computation is an ordered pipeline (Figure 4). (1) Trend: an OLS line "
         "T_t = b*t + a is fitted not on raw samples but on per-period block means at "
         "their block centres, so averaging a full season cancels the periodic part and "
         "it cannot leak into the slope. (2) Seasonal: the detrended series d_t = y_t - "
         "T_t is averaged by phase j = t mod m and re-centred so the indices S_j sum to "
         "zero. (3) Residual + AR(1): r_t = d_t - S_(t mod m); rho is the residuals' lag-1 "
         "autocorrelation clamped to [0, 0.95]. (4) Point forecast at horizon h: "
         "y_hat = T_(n-1+h) + S_((n-1+h) mod m) + rho^h * r_(n-1). (5) Interval: an "
         "expanding-window one-step backtest yields RMSE, giving the 95% band "
         "y_hat +/- 1.96*RMSE. The optimizer then takes the forecast peak: (6) the robust "
         "peak is the nearest-rank p95; (7) load in resource units L = (p95/100)*u_cur*"
         "gamma; (8) it solves min u s.t. L <= tau*u over u in {1..u_cur}; (9) the exact "
         "solution is u* = min(u_cur, ceil(L/tau)), repeated for memory in 256 MB blocks, "
         "yielding the cost saving and the snapped GCP machine type. Every term in "
         "Figure 4 is the exact expression evaluated in core.forecaster and core.optimizer.",
         "전 과정은 순서가 있는 파이프라인이다(그림 4). (1) 추세: OLS 직선 T_t = b·t + a를 원시 "
         "표본이 아니라 주기별 블록 평균(블록 중심 위치)에 적합한다. 한 주기를 평균하면 계절 성분이 "
         "상쇄되어 기울기에 누설되지 않는다. (2) 계절: 탈추세 계열 d_t = y_t − T_t를 위상 "
         "j = t mod m별로 평균하고, 지수 S_j의 합이 0이 되도록 재중심화한다. (3) 잔차 + AR(1): "
         "r_t = d_t − S_(t mod m), ρ는 잔차의 시차-1 자기상관으로 [0, 0.95]에 클램프한다. "
         "(4) 점예측(지평 h): ŷ = T_(n-1+h) + S_((n-1+h) mod m) + ρ^h·r_(n-1). (5) 구간: "
         "확장창 1-스텝 백테스트로 RMSE를 구해 95% 구간 ŷ ± 1.96·RMSE를 만든다. 이어 옵티마이저는 "
         "예측 피크를 입력으로 받는다. (6) 강건 피크는 근접순위 p95, (7) 자원단위 부하 "
         "L = (p95/100)·u_cur·γ, (8) u ∈ {1..u_cur}에서 min u s.t. L ≤ τ·u를 풀고, (9) 정확 해 "
         "u* = min(u_cur, ceil(L/τ))이며 메모리는 256MB 블록으로 동일하게 계산해 비용 절감과 GCP "
         "머신 타입 스냅을 산출한다. 그림 4의 모든 항은 core.forecaster·core.optimizer가 실제로 "
         "계산하는 식이다."),
        ("IMG", ("method_math",
                 "Figure 4. End-to-end computation pipeline with the exact equations "
                 "(STL + AR(1) forecast -> 95% interval; p95 peak -> integer program -> "
                 "machine type).",
                 "그림 4. 정확한 수식이 포함된 전 과정 연산 파이프라인 (STL + AR(1) 예측 → 95% 구간; "
                 "p95 피크 → 정수계획 → 머신 타입).", 6.6)),
        ("H2", "3.2. Decomposition and AR residual correction", "3.2. 계절-추세 분해와 AR 잔차 보정"),
        ("B",
         "The series is modelled as trend + seasonal + residual. Trend is fitted by least "
         "squares over per-period block means, so averaging a full season cancels the "
         "periodic component and it cannot leak into the slope; seasonal indices are the "
         "mean detrended value per phase, re-centred to sum to zero. Real CPU is strongly "
         "autocorrelated (measured lag-1 0.5-0.9), so a damped AR(1) residual correction "
         "is added: y_hat = (trend + seasonal) + rho^h * (last actual - last fitted), where "
         "rho is the residuals' lag-1 autocorrelation and h the horizon. This term makes "
         "the model beat both baselines.",
         "예측기는 시계열을 추세+계절+잔차로 분해한다. 추세는 주기별 블록 평균에 최소제곱으로 추정해 "
         "계절 성분의 추세 누설을 제거하고, 계절 지수는 위상별 평균 탈추세 값을 0합으로 재중심화한다. "
         "실제 CPU는 시간 간 강하게 자기상관되므로(측정된 시차-1 자기상관 0.5–0.9), 분해 예측에 AR(1) "
         "잔차 보정을 더한다: ŷ = (추세+계절) + ρ^h·(직전 실측값 − 직전 적합값). ρ는 잔차의 시차-1 "
         "자기상관, h는 예측 지평이다. 이 항이 모델을 두 기준선보다 우수하게 만든다."),
        ("H2", "3.3. Point forecast and 95% prediction interval", "3.3. 점예측과 95% 예측구간"),
        ("B",
         "Each forecast has (i) a point estimate and (ii) a 95% prediction interval - the "
         "range the actual future value should fall in with 95% probability. The interval "
         "is sized from out-of-sample one-step backtest error (RMSE): "
         "[y_hat - 1.96*RMSE, y_hat + 1.96*RMSE], so uncertainty is not understated on "
         "short series. Measured coverage (PICP) is 0.93-0.98 versus the nominal 0.95, so "
         "the interval is well-calibrated (Figure 8). Unlike a black box, exposing "
         "'forecast X% +/- range' makes risk legible (white-box). The engine uses only the "
         "Python standard library.",
         "각 예측은 (i) 점예측치와 (ii) 95% 예측구간(실제 미래값이 95% 확률로 들 범위)으로 구성된다. "
         "구간은 표본 외 1-스텝 백테스트 오차(RMSE)로 산정한다: [예측치 − 1.96·RMSE, 예측치 + "
         "1.96·RMSE]. 짧은 시계열에서 불확실성을 과소평가하지 않기 위함이다. 측정된 커버리지(PICP)는 "
         "0.93–0.98로 공칭 0.95와 거의 일치해 잘 보정된다(그림 8). 블랙박스와 달리 '예측 X% ± 범위'를 "
         "그대로 노출해 위험을 가늠하게 한다(화이트박스). 엔진은 Python 표준 라이브러리만 사용한다."),

        ("H1", "4. Integer-Programming Resizing and Real Application",
         "4. 정수계획 리사이징 및 실제 적용"),
        ("H2", "4.1. Robust peak (p95) and the headroom constraint",
         "4.1. 강건 피크(p95)와 헤드룸 제약"),
        ("B",
         "For sizing, the 'peak' is the 95th percentile, not the maximum. The maximum is "
         "vulnerable to a single transient spike that would force permanent over-"
         "provisioning; p95 discards the top ~5% so rare spikes do not drive sizing - a "
         "robust peak. Conversely, frequently recurring load (e.g., a daily batch spike) "
         "exceeds the 5% threshold and is included in p95, so that host is correctly held "
         "(the bursty demo host stays unchanged).",
         "사이징 기준 '피크'로 최댓값이 아니라 95퍼센타일(p95)을 쓴다. 최댓값은 단발 스파이크 1점에 "
         "취약해 영구 과프로비저닝을 유발하지만, p95는 극단 ~5%를 제외해 드문 급등에 흔들리지 않는다 — "
         "강건(robust) 피크다. 반대로 자주 반복되는 부하(예: 매일 배치 스파이크)는 빈도가 5%를 넘어 "
         "p95에 포함되므로 그 호스트는 올바르게 유지된다(데모의 버스트 호스트가 0% 유지)."),
        ("B",
         "The engine picks the smallest integer allocation keeping the robust-peak "
         "utilisation under target while preserving a margin: p95_peak * safety_margin <= "
         "target_utilisation * allocation. safety_margin (e.g. 1.2) absorbs forecast error; "
         "target_utilisation (e.g. 0.65) is the steady-state ceiling. The bounded space "
         "(vCPU in [1, current], memory in 256 MB blocks) is solved by exact enumeration; "
         "vCPU and memory are optimised independently.",
         "엔진은 안전 마진을 유지하며 강건 피크 가동률을 목표 이하로 두는 최소 정수 할당을 선택한다: "
         "p95_peak × safety_margin ≤ target_utilisation × allocation. safety_margin(예 1.2)은 "
         "예측 오차 버퍼, target_utilisation(예 0.65)은 정상상태 상한이다. 유계 공간(vCPU∈[1,현재], "
         "메모리 256MB 블록)을 전수 열거로 풀며 vCPU·메모리를 독립 최적화한다."),
        ("H2", "4.2. GCP machine-type mapping and budget-guarded real resize",
         "4.2. GCP 머신 타입 매핑과 예산 가드 실제 리사이즈"),
        ("B",
         "Abstract (vcpu, memory) recommendations snap to predefined GCP instances "
         "(E2/N2/C2/C3). Real GCE hosts are physically resized via the Compute API, but a "
         "machine type whose monthly cost exceeds the budget (e.g. 300,000 KRW/month) is "
         "rejected, and protected instances (e.g. the work machine) are excluded.",
         "추상적 (vcpu, memory) 권장값은 사전정의 GCP 인스턴스(E2/N2/C2/C3)로 스냅된다. 실제 GCE "
         "호스트는 Compute API로 물리적으로 리사이즈하되, 월 비용이 한도(예: 300,000원/월)를 초과하는 "
         "타입은 거부하고 보호 인스턴스(작업 머신 등)는 제외한다."),

        ("H1", "5. Implementation and Deployment", "5. 구현 및 배포"),
        ("B",
         "The backend is FastAPI with SQLAlchemy 2.0 async; the frontend is React 19 built "
         "with Vite and served by nginx; both run as non-root containers on Cloud Run. The "
         "dashboard's top menu offers Dashboard / History (full metric tables + activity "
         "log) / Test Scenarios (forecast, optimise, real GCP sync, real resize, model "
         "evaluation). Every forecast and resize is persisted to an audit-log table. The "
         "real fleet runs diurnal load generators so Cloud Monitoring records a "
         "representative live signal.",
         "백엔드는 SQLAlchemy 2.0 async의 FastAPI, 프론트엔드는 Vite 빌드 React 19(nginx 서빙)이며 둘 "
         "다 비루트 컨테이너로 Cloud Run에 배포된다. 대시보드 상단 메뉴: Dashboard / History(전체 "
         "메트릭·활동 로그) / Test Scenarios(예측·최적화·실제 GCP 동기화·실제 리사이즈·모델 평가). "
         "모든 예측·리사이즈는 감사 로그에 영속 기록된다. 실제 플릿은 일주기 부하 생성기를 구동해 "
         "Cloud Monitoring에 대표적인 실측 신호를 남긴다."),
        ("URL", f"Live dashboard: {FE}", f"라이브 대시보드: {FE}"),
        ("URL", f"Live API (Swagger /docs): {BE}", f"라이브 API (Swagger /docs): {BE}"),

        ("H1", "6. Statistical Representativeness of the Test Data",
         "6. 시험 데이터의 통계적 대표성"),
        ("B",
         "Demo/seed metrics are grounded in public traces: Azure Resource Central "
         "(SOSP'17; 60% of VMs average < 20% CPU; interactive vs delay-insensitive), "
         "Barroso & Holzle (servers mostly 10-50%), Alibaba 2018 (batch ~29.3%, service "
         "~7.4%). Six archetypes are generated over 14 days hourly, not as a repeating "
         "curve but with AR(1) autocorrelation, day-to-day amplitude variation, a slow "
         "trend, heteroscedastic noise, and rare anomalies. The generated coefficient of "
         "variation is 0.28-1.11 and lag-1 autocorrelation 0.5-0.9, matching real "
         "distributions (versus ~0 for uniform synthetic data).",
         "데모/시드 메트릭은 공개 트레이스에 근거한다. Azure Resource Central(SOSP'17): VM 60%가 평균 "
         "CPU<20%, 대화형/지연무관 분류. Barroso & Hölzle: 서버 대부분 10–50% 대역. Alibaba 2018: "
         "배치 평균 29.3%, 서비스 7.4%. 6개 아키타입을 14일×시간단위로 생성하되 단순 반복 곡선이 아니라 "
         "AR(1) 자기상관·날짜별 진폭 변동·완만한 추세·이분산 잡음·희귀 이상치를 가산한다. 생성 데이터의 "
         "변동계수(CV)는 0.28–1.11, 시차-1 자기상관은 0.5–0.9로 실측 분포와 일치한다(균일 합성 데이터의 "
         "CV≈0과 대조)."),

        ("H1", "7. Model Evaluation (Paper-Level)", "7. 모델 평가 (논문 수준)"),
        ("B",
         "A single MAPE on uniform data cannot prove validity. The model is compared "
         "against two standard baselines (naive = last value, seasonal-naive = value one "
         "period ago) on an expanding-window one-step backtest, with academic statistics "
         "MASE (Hyndman & Koehler 2006), sMAPE, RMSE; interval calibration via PICP "
         "(coverage) and MPIW (width); and the Diebold-Mariano test (1995) for statistical "
         "significance.",
         "균일 데이터의 단일 MAPE는 유효성을 입증하지 못한다. 확장 윈도우 1-스텝 백테스트로 두 표준 "
         "기준선(naive=직전값, seasonal-naive=한 주기 전 값)과 비교하고, 학술 지표 MASE(Hyndman & "
         "Koehler 2006)·sMAPE·RMSE, 구간 보정 PICP(커버리지)·MPIW(폭), Diebold-Mariano 검정(1995)으로 "
         "유의성을 평가한다."),
        ("B",
         "Results: the model beats seasonal-naive (RMSE) on all 6 archetypes, MASE < 1 on "
         "5/6, and is significantly better (Diebold-Mariano p<0.05) on 4/6. Interval "
         "coverage is 0.93-0.98 (nominal 0.95). For flat/idle workloads the last-value "
         "baseline is strong, so significance is lower there - an honest property of those "
         "signals.",
         "결과: 모델은 6/6에서 seasonal-naive를 RMSE로 능가, 5/6에서 MASE<1, 4/6에서 "
         "Diebold-Mariano p<0.05로 유의하게 우수하다. 구간 커버리지는 0.93–0.98(공칭 0.95)로 잘 "
         "보정된다. 저활용 워크로드는 직전값 기준선이 강해 유의성이 낮은데, 이는 평탄·준유휴 신호의 "
         "본질적 특성으로 정직하게 보고한다."),
        ("IMG", ("fig_forecast_overlay",
                 "Figure 5. One-step forecast vs actual with the 95% interval.",
                 "그림 5. 1-스텝 예측 vs 실측과 95% 예측구간.", 5.6)),
        ("IMG", ("fig_error_rmse",
                 "Figure 6. Forecast error (RMSE): model vs baselines (lower is better).",
                 "그림 6. 예측 오차(RMSE): 모델 vs 기준선(낮을수록 우수).", 6.0)),
        ("IMG", ("fig_mase",
                 "Figure 7. MASE - below 1 beats seasonal-naive.",
                 "그림 7. MASE — 1 미만이면 seasonal-naive 능가.", 6.0)),
        ("IMG", ("fig_coverage",
                 "Figure 8. 95% prediction-interval calibration (PICP) vs nominal 0.95.",
                 "그림 8. 95% 예측구간 보정(PICP) — 공칭 0.95 대비.", 6.0)),
        ("IMG", ("fig_residual_acf",
                 "Figure 9. Residual autocorrelation - within the white-noise band = good fit.",
                 "그림 9. 잔차 자기상관 — 백색잡음 대역 내면 적합 양호.", 6.0)),

        ("H1", "8. Related Work and Differentiation", "8. 관련 연구 및 차별화"),
        ("B",
         "The market splits into (a) Kubernetes/cloud SaaS optimizers (CAST AI, Sedai, "
         "StormForge - telemetry egress, K8s-assumed), (b) cloud-provider recommenders "
         "(AWS Compute Optimizer, Azure Advisor - single cloud, short windows), and (c) "
         "on-prem monitors (SolarWinds, ManageEngine - threshold-based descriptive "
         "reporting). MetricLens occupies the gap: (1) air-gapped and self-contained, "
         "(2) GPU-free and lightweight, (3) prescriptive SLO-constrained integer "
         "programming, (4) white-box explainability, (5) infrastructure-agnostic "
         "(generic VM/bare-metal).",
         "시장은 (a) K8s/클라우드 SaaS 최적화기(CAST AI·Sedai·StormForge — 텔레메트리 외부 전송·K8s "
         "전제), (b) CSP 추천기(AWS Compute Optimizer·Azure Advisor — 단일 클라우드·짧은 윈도우), "
         "(c) 온프레 모니터(SolarWinds·ManageEngine — 임계치 기반 서술적 리포팅)로 나뉜다. MetricLens는 "
         "그 사이의 공백을 차지한다: (1) 에어갭·자립형, (2) GPU-프리 경량, (3) SLO 제약 처방적 정수계획, "
         "(4) 화이트박스 설명가능성, (5) 인프라 비종속(일반 VM/베어메탈)."),

        ("H1", "9. Conclusion", "9. 결론"),
        ("B",
         "MetricLens AI delivers accurate CPU-only forecasting and exact integer-"
         "programming resizing as a dependency-free, cloud-native service, monitors and "
         "resizes real GCP instances, and proves model validity with a paper-level "
         "evaluation (baselines, MASE, Diebold-Mariano, interval calibration). By turning "
         "passive telemetry into quantitative, SLO-aware sizing guidance, it offers a "
         "practical path to cutting infrastructure cost and energy waste on-premise.",
         "MetricLens AI는 정확한 CPU-온리 예측과 정수계획 리사이징을 의존성 없는 클라우드 네이티브 "
         "서비스로 제공하고, 실제 GCP 인스턴스를 모니터링·리사이즈하며, 모델 유효성을 논문 수준의 "
         "평가(기준선·MASE·Diebold-Mariano·구간 보정)로 입증한다. 수동적 텔레메트리를 정량적·SLO 인지 "
         "사이징 지침으로 전환해, 온프레미스 비용·전력 낭비를 줄이는 실용적 경로를 제시한다."),
    ]


def kfont(p):
    for r in p.runs:
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rf)
        rf.set(qn("w:eastAsia"), "Malgun Gothic")


def build(lang):
    doc = Document(str(TEMPLATE))
    items = sections(lang)
    idx = 0 if lang == "en" else 1  # which language string

    def txt(t):
        return t[1] if lang == "en" else t[2]

    paras = doc.paragraphs
    title, auth, affil = items[0], items[1], items[2]

    def set_p(p, s):
        for r in list(p.runs):
            r.text = ""
        (p.runs[0] if p.runs else p.add_run("")).text = s
        if not p.runs:
            p.add_run(s)
        kfont(p)

    set_p(paras[0], txt(title)); set_p(paras[1], txt(auth)); set_p(paras[2], txt(affil))
    for p in paras[3:]:
        p._element.getparent().remove(p._element)

    for kind, *rest in items[3:]:
        if kind in ("H1", "H2", "B", "URL"):
            style = {"H1": "Header 1", "H2": "Header 2", "B": "Body", "URL": "Body"}[kind]
            kfont(doc.add_paragraph(rest[0] if lang == "en" else rest[1], style=style))
        elif kind == "IMG":
            name, cap_en, cap_kr, width = rest[0]
            path = fig(name, lang)
            if path.exists():
                doc.add_picture(str(path), width=Inches(width))
                kfont(doc.add_paragraph(cap_en if lang == "en" else cap_kr, style="Body"))

    out = ROOT / "docs" / ("development_report_usenix.docx" if lang == "en"
                           else "development_report_usenix_kr.docx")
    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    build("en")   # English primary
    build("kr")   # Korean variant
